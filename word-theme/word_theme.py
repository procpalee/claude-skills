# -*- coding: utf-8 -*-
"""
python-docx 워드 테마 헬퍼.

워드(.docx)를 만들 때 word_themes.py(WORD_THEMES) 기반의 일관된
제목부/제목스타일/표/머리·바닥글 서식을 코드로 적용한다.
(워드 전용 테마에 없는 이름은 ../excel-theme/excel_theme.py 팔레트로 폴백)

제목부 스타일 5종(bar/side/band/center/meta)과 문서유형 프리셋(doc_type)은
word_themes.py 의 TITLE_STYLES / DOC_TYPES 참조.

전형적 사용:
    from word_theme import new_doc, add_title, add_heading, add_paragraph, add_table

    doc = new_doc()                              # 기본 default (모노톤 그레이)
    doc = new_doc("procpa", doc_type="report")   # 브랜드 블루 + 보고서 프리셋
    add_title(doc, "FY25 손상평가 보고서", "OO사 · 2026.03",
              meta={"작성": "홍길동", "일자": "2026-03-15"})
    add_heading(doc, "1. 평가 개요", level=1)
    add_paragraph(doc, "본 보고서는 ...")
    add_table(doc, headers=[...], rows=[...], number_cols={2: "accounting"})
    doc.save("out.docx")

new_doc 이 테마를 문서에 기억시키므로 이후 헬퍼에 theme= 를 반복할 필요가 없다
(명시하면 그 값이 우선). 색을 바꾸려면 word_themes.py 의 WORD_THEMES 를 수정한다.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 팔레트 폴백용으로 excel-theme 의 팔레트를 참조한다.
import os as _os
import sys as _sys
_EXCEL_DIR = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))),
                           "excel-theme")
if _EXCEL_DIR not in _sys.path:
    _sys.path.insert(0, _EXCEL_DIR)
from excel_theme import PALETTES as _EXCEL_PALETTES, get_palette as _excel_get_palette  # noqa: E402

# 워드 전용 테마 레지스트리(엑셀과 색 일치 불필요). 기본 = default.
_THIS_DIR = _os.path.dirname(_os.path.abspath(__file__))
if _THIS_DIR not in _sys.path:
    _sys.path.insert(0, _THIS_DIR)
from word_themes import (WORD_THEMES, WORD_ALIASES, DEFAULT_WORD_THEME,  # noqa: E402
                         TITLE_STYLES, DOC_TYPES)

THEMES = {**_EXCEL_PALETTES, **WORD_THEMES}
DEFAULT_THEME = DEFAULT_WORD_THEME   # "default"


def get_theme(name=None):
    """워드 전용 테마 우선 조회(별칭 해석) → 없으면 엑셀 팔레트로 폴백. dict 는 그대로 통과."""
    if name is None:
        name = DEFAULT_THEME
    if isinstance(name, dict):
        return name
    name = WORD_ALIASES.get(name, name)   # 구 이름(workpaper 등) → 신 이름
    if name in WORD_THEMES:
        return WORD_THEMES[name]
    return _excel_get_palette(name)


def _resolve_theme(doc, theme):
    """명시 theme > new_doc 에 기억된 테마 > 기본 테마 순으로 해석."""
    if theme is not None:
        return theme if isinstance(theme, dict) else get_theme(theme)
    remembered = getattr(doc, "_wt_theme", None)
    return remembered if remembered is not None else get_theme(None)


__all__ = [
    "new_doc", "apply_base_style", "add_title", "add_heading",
    "add_paragraph", "add_bullets", "add_callout", "add_table",
    "add_page_number_footer", "THEMES", "DEFAULT_THEME",
    "TITLE_STYLES", "DOC_TYPES",
]


# ──────────────────────────────────────────────────────────────
# 저수준 유틸
# ──────────────────────────────────────────────────────────────
def _rgb(hex6):
    return RGBColor.from_string(hex6)


def _set_run_font(run, name, size=None, color=None, bold=None):
    """run 에 라틴+동아시아 글꼴, 크기, 색, 굵기 적용."""
    run.font.name = name
    # 동아시아(한글) 글꼴은 별도 지정해야 적용된다
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.font.bold = bold


def _char_spacing(run, pt):
    """자간(문자 간격, pt). 캡션 라벨 등에 사용."""
    rpr = run._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(int(pt * 20)))
    rpr.append(sp)


def _shade_cell(cell, hex6):
    """표 셀 배경색."""
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcpr.append(shd)


def _set_cell_borders(cell, edges):
    """셀 테두리. edges: {'top': (size8, hex), 'bottom':..., 'left':..., 'right':...}
    size8 은 1/8pt 단위 (예: 8 = 1pt, 16 = 2pt)."""
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge, (sz, col) in edges.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), col)


def _set_table_width(tbl, pct=None):
    """표 너비를 페이지 폭 대비 % 로 지정 (pct=100 → 전폭)."""
    tblPr = tbl._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), str(int(pct * 50)))    # pct 단위 = 1/50 %
    tblPr.append(tblW)


def _enable_table_autofit(tbl, mode="content"):
    """표 자동맞춤. tblLayout=autofit + 셀 고정폭(tcW) 제거.
    mode="content" : 표 너비 auto — 열을 내용폭에 맞게(표가 좁아질 수 있음).
    mode="window"  : 표 너비 100%(pct 5000) — 페이지 폭을 채우고 열은 내용 비율로 분배."""
    tbl.autofit = True
    tbl.allow_autofit = True
    tblPr = tbl._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    if mode == "window":
        tblW.set(qn("w:type"), "pct"); tblW.set(qn("w:w"), "5000")  # 5000 = 100%
    else:
        tblW.set(qn("w:type"), "auto"); tblW.set(qn("w:w"), "0")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)
    # 셀에 박힌 고정폭 제거 (autofit 방해 요소)
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcPr.remove(tcW)


def _tighten_cell(cell):
    """셀 안 문단의 위·아래 여백과 줄간격을 줄여 행 높이를 내용에 맞춘다."""
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.0


def _bar_paragraph(doc, hex6, height_pt=3, before=2, after=6, line_pt=None):
    """가는 색 막대(가로줄)를 문단 하단 테두리로 표현 — 타이틀 강조 바.
    line_pt 지정 시 빈 문단의 줄높이를 고정(pt)해 위 요소와의 간격을 최소화한다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if line_pt is not None:
        p.paragraph_format.line_spacing = Pt(line_pt)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(height_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex6)
    pbdr.append(bottom)
    ppr.append(pbdr)
    return p


def _tighten_title_lines(p, size):
    """제목이 두 줄로 꺾일 때 줄간격이 벌어지지 않도록 고정 줄높이 적용
    (맑은 고딕 계열은 자연 줄높이가 ~1.34em 로 커서 큰 제목에서 벌어져 보인다)."""
    p.paragraph_format.line_spacing = Pt(round(size * 1.15))


def _fmt_number(val, fmt=None):
    """엑셀 NUMBER_FORMATS 규약을 워드 텍스트로 재현.
    음수 = 괄호(색은 호출부에서 negative 적용), 0 = 대시."""
    fmt = str(fmt or "accounting")
    if val == 0:
        return "-"
    if "percent" in fmt:
        s = f"{abs(val):.2%}" if "acct" in fmt else f"{abs(val):.1%}"
    elif "decimal" in fmt:
        s = f"{abs(val):,.2f}"
    elif "multiple" in fmt:
        s = f"{abs(val):,.1f}x"
    else:                       # accounting / thousands / 기타
        s = f"{abs(val):,.0f}"
    return f"({s})" if val < 0 else s


# ──────────────────────────────────────────────────────────────
# 기본 스타일 · 문서 생성
# ──────────────────────────────────────────────────────────────
def apply_base_style(doc, theme=None):
    """Normal/제목 스타일에 팔레트 글꼴·색 적용."""
    th = _resolve_theme(doc, theme)
    c, s, font = th["colors"], th["sizes"], th["font_name"]

    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(s["body"])
    normal.font.color.rgb = _rgb(c["text"])
    # Normal 의 동아시아 글꼴
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), font)

    # 내장 Heading 스타일 색/크기/글꼴 재정의
    heading_map = {
        "Heading 1": (s["h1"], c["primary"]),
        "Heading 2": (s["h2"], c["primary_dark"]),
        "Heading 3": (s["h3"], c["secondary"]),
    }
    for name, (size, col) in heading_map.items():
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = font
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = _rgb(col)
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            rfonts.set(qn(attr), font)
    return doc


def new_doc(theme=None, doc_type=None):
    """팔레트 기본 스타일이 적용된 새 문서 생성.

    theme    : WORD_THEMES 키(default/procpa) 또는 dict. 문서에 기억되어
               이후 헬퍼들이 자동 상속한다(헬퍼에 theme= 명시 시 그 값 우선).
    doc_type : DOC_TYPES 키(report/memo/proposal/opinion/workpaper).
               add_title 의 기본 제목 스타일·캡션이 결정된다.
    """
    doc = Document()
    th = get_theme(theme)
    doc._wt_theme = th
    doc._wt_doc_type = doc_type
    apply_base_style(doc, th)
    add_page_number_footer(doc, th)
    return doc


# ──────────────────────────────────────────────────────────────
# 제목부 (6종 스타일)
# ──────────────────────────────────────────────────────────────
def add_title(doc, title, subtitle=None, theme=None, style=None, meta=None):
    """문서 제목부. style 미지정 시 new_doc(doc_type=) 프리셋 → 기본 "bar".

    style : "bar"    제목 + 하단 강조 바 (기본 — 일반 보고서)
            "side"   좌측 세로 악센트 바 + 제목/부제           — 간단한 메모·검토
            "band"   전폭 색 밴드(짙은 배경 + 흰 제목)         — 제안서·트렌디 산출물
            "center" 중앙 정렬 + 상하 가는 선                  — 공식 의견서·공문
            "meta"   제목(좌) + 문서정보 표(우) + 하단 굵은선  — 조서
    meta  : {"caption": 문서유형 라벨(band), "작성": "홍길동", ...} —
            caption 외 항목은 center 의 메타 라인 또는 meta 의 문서정보 표에 표기된다.
    """
    th = _resolve_theme(doc, theme)
    preset = DOC_TYPES.get(getattr(doc, "_wt_doc_type", None) or "", {})
    style = style or preset.get("title_style") or "bar"
    if style not in TITLE_STYLES:
        raise KeyError(f"알 수 없는 제목 스타일 '{style}'. 사용 가능: {', '.join(TITLE_STYLES)}")
    meta = dict(meta or {})
    caption = meta.pop("caption", None) or preset.get("caption")

    builder = {"bar": _title_bar, "side": _title_side,
               "band": _title_band, "center": _title_center, "meta": _title_meta}[style]
    builder(doc, th, title, subtitle, caption, meta)
    return doc


def _title_text(container, th, text, *, size, color, bold, align=None,
                before=0, after=2):
    """제목부 내부 공용: 문단 하나 추가."""
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, th["font_name"], size=size, color=color, bold=bold)
    return p


def _meta_line(meta, sep=" · "):
    return sep.join(f"{k} {v}" for k, v in meta.items())


def _title_bar(doc, th, title, subtitle, caption, meta):
    """A. 제목 + (선택) 부제 + 강조 바 — 기본."""
    c, s = th["colors"], th["sizes"]
    p = _title_text(doc, th, title, size=s["title"], color=c["title_color"], bold=True)
    _tighten_title_lines(p, s["title"])
    if subtitle:
        _title_text(doc, th, subtitle, size=s["subtitle"], color=c["secondary"], bold=False)
    _bar_paragraph(doc, c["primary"], height_pt=2.5)


def _title_side(doc, th, title, subtitle, caption, meta):
    """C. 좌측 세로 악센트 바 + 제목/부제 — 간단한 메모·검토."""
    c, s = th["colors"], th["sizes"]
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_borders(cell, {"left": (28, c["primary"])})
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _tighten_title_lines(p, s["title"] - 4)
    run = p.add_run(title)
    _set_run_font(run, th["font_name"], size=s["title"] - 4, color=c["title_color"], bold=True)
    if subtitle:
        sp = cell.add_paragraph()
        sp.paragraph_format.left_indent = Pt(10)
        srun = sp.add_run(subtitle)
        _set_run_font(srun, th["font_name"], size=s["subtitle"], color=c["secondary"], bold=False)
    doc.add_paragraph()


def _title_band(doc, th, title, subtitle, caption, meta):
    """D. 전폭 색 밴드(짙은 배경 + 흰 제목) — 제안서·표지."""
    c, s = th["colors"], th["sizes"]
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_width(tbl, pct=100)
    cell = tbl.cell(0, 0)
    _shade_cell(cell, c["primary"])
    _set_cell_borders(cell, {e: (4, c["primary"])
                             for e in ("top", "bottom", "left", "right")})
    first = cell.paragraphs[0]
    first.paragraph_format.space_before = Pt(14)
    first.paragraph_format.left_indent = Pt(12)
    if caption:
        run = first.add_run(caption)
        _set_run_font(run, th["font_name"], size=s["small"], color=c["secondary_lt"], bold=False)
        _char_spacing(run, 2)
        tp = cell.add_paragraph()
    else:
        tp = first
    tp.paragraph_format.left_indent = Pt(12)
    tp.paragraph_format.space_after = Pt(2)
    _tighten_title_lines(tp, s["title"] - 2)
    run = tp.add_run(title)
    _set_run_font(run, th["font_name"], size=s["title"] - 2, color="FFFFFF", bold=True)
    last = tp
    if subtitle:
        sp = cell.add_paragraph()
        sp.paragraph_format.left_indent = Pt(12)
        srun = sp.add_run(subtitle)
        _set_run_font(srun, th["font_name"], size=s["subtitle"], color=c["secondary_lt"], bold=False)
        last = sp
    last.paragraph_format.space_after = Pt(14)
    doc.add_paragraph()


def _title_center(doc, th, title, subtitle, caption, meta):
    """E. 중앙 정렬 + 상하 가는 선 — 공식 의견서·공문·회신문."""
    c, s = th["colors"], th["sizes"]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(tbl, pct=88)
    cell = tbl.cell(0, 0)
    _set_cell_borders(cell, {"top": (4, c["secondary_lt"]),
                             "bottom": (4, c["secondary_lt"])})
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _tighten_title_lines(p, s["title"] - 4)
    run = p.add_run(title)
    _set_run_font(run, th["font_name"], size=s["title"] - 4, color=c["title_color"], bold=True)
    last = p
    if subtitle:
        sp = cell.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sp.add_run(subtitle)
        _set_run_font(srun, th["font_name"], size=s["subtitle"], color=c["secondary"], bold=False)
        last = sp
    last.paragraph_format.space_after = Pt(10)
    if meta:
        _title_text(doc, th, _meta_line(meta), size=s["small"], color=c["secondary_lt"],
                    bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6)
    else:
        doc.add_paragraph()


def _title_meta(doc, th, title, subtitle, caption, meta):
    """F. 제목(좌) + 문서정보 표(우) + 하단 굵은선 — 조서·품질관리 문서."""
    c, s = th["colors"], th["sizes"]
    outer = doc.add_table(rows=1, cols=2)
    outer.autofit = False
    left, right = outer.cell(0, 0), outer.cell(0, 1)
    left.width, right.width = Cm(10.5), Cm(5.5)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    _tighten_title_lines(p, s["title"] - 5)
    run = p.add_run(title)
    _set_run_font(run, th["font_name"], size=s["title"] - 5, color=c["title_color"], bold=True)
    if subtitle:
        sp = left.add_paragraph()
        srun = sp.add_run(subtitle)
        _set_run_font(srun, th["font_name"], size=s["subtitle"], color=c["secondary"], bold=False)

    if meta:
        inner = right.add_table(rows=len(meta), cols=2)
        inner.alignment = WD_TABLE_ALIGNMENT.RIGHT
        # 셀 삽입으로 생긴 선행 빈 문단 제거
        lead = right.paragraphs[0]
        lead._p.getparent().remove(lead._p)
        hdr_fill = c.get("table_header_fill", c["band"])
        hdr_text = c.get("table_header_text", c["text"])
        for i, (k, v) in enumerate(meta.items()):
            kc, vc = inner.cell(i, 0), inner.cell(i, 1)
            kc.width, vc.width = Cm(2.0), Cm(3.2)
            _shade_cell(kc, hdr_fill)
            for cc, txt, bold in ((kc, str(k), True), (vc, str(v), False)):
                _set_cell_borders(cc, {e: (4, c["border_in"])
                                       for e in ("top", "bottom", "left", "right")})
                cp = cc.paragraphs[0]
                crun = cp.add_run(txt)
                _set_run_font(crun, th["font_name"], size=s["small"],
                              color=hdr_text if bold else c["text"], bold=bold)
            _tighten_cell(kc); _tighten_cell(vc)

    # 하단 선: 빈 문단 줄높이를 고정(2pt)해 제목/표와의 간격을 최소화
    _bar_paragraph(doc, c["primary"], height_pt=2, before=2, line_pt=2)


# ──────────────────────────────────────────────────────────────
# 콘텐츠 헬퍼
# ──────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, theme=None):
    """팔레트 색이 입혀진 제목. (apply_base_style 가 색을 보장)"""
    th = _resolve_theme(doc, theme)
    c, s = th["colors"], th["sizes"]
    col = {1: c["primary"], 2: c["primary_dark"], 3: c["secondary"]}.get(level, c["text"])
    size = {1: s["h1"], 2: s["h2"], 3: s["h3"]}.get(level, s["body"])
    p = doc.add_paragraph(style=f"Heading {level}" if level <= 3 else None)
    run = p.add_run(text)
    _set_run_font(run, th["font_name"], size=size, color=col, bold=True)
    return p


def add_paragraph(doc, text, theme=None, bold=False, color=None):
    th = _resolve_theme(doc, theme)
    c, s = th["colors"], th["sizes"]
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, th["font_name"], size=s["body"],
                  color=color or c["text"], bold=bold)
    return p


def add_bullets(doc, items, theme=None, numbered=False):
    """글머리표/번호 목록. items 원소가 (level, text) 면 들여쓰기(level 0/1)."""
    th = _resolve_theme(doc, theme)
    c, s = th["colors"], th["sizes"]
    base = "List Number" if numbered else "List Bullet"
    out = []
    for it in items:
        lvl, txt = it if isinstance(it, tuple) else (0, it)
        style = base if lvl == 0 else "%s %d" % (base, lvl + 1)
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph(style=base)
        run = p.add_run(txt)
        _set_run_font(run, th["font_name"], size=s["body"], color=c["text"])
        out.append(p)
    return out


def add_callout(doc, text, theme=None):
    """note 색 배경 + accent 좌측선 강조 박스(1x1 표로 구현)."""
    th = _resolve_theme(doc, theme)
    c, s = th["colors"], th["sizes"]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _shade_cell(cell, c["note"])
    _set_cell_borders(cell, {"left": (24, c["accent"]),
                             "top": (4, c["note"]), "bottom": (4, c["note"]),
                             "right": (4, c["note"])})
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(text)
    _set_run_font(run, th["font_name"], size=s["body"], color=c["text"], bold=False)
    doc.add_paragraph()
    return tbl


def add_table(doc, headers, rows, theme=None,
              number_cols=None, total_last_row=False, autofit=True):
    """헤더+데이터 표를 팔레트 서식으로 생성.

    headers       : 헤더 셀 리스트
    rows          : [[...], ...] 데이터 행
    number_cols   : {열index(0-based): 포맷문자열} — 우측정렬 + 회계 표기.
                    숫자는 엑셀 NUMBER_FORMATS 규약을 따른다:
                    음수 = negative 색 괄호, 0 = 대시. (percent/decimal/multiple 지원)
    total_last_row: 마지막 행을 합계 행으로 강조
    autofit       : 표 너비 모드. 모두 행 높이는 내용에 맞게 압축한다.
                    True/"content"(기본) — 열을 내용폭에 맞춤(표가 좁아질 수 있음).
                    "window"             — 페이지 폭 100%로 채우고 열은 내용 비율 분배.
                    False                — 고정폭 균등 표(자동맞춤 끔).
    """
    th = _resolve_theme(doc, theme)
    c, s = th["colors"], th["sizes"]
    number_cols = number_cols or {}

    # 테마 옵션(있으면 사용, 없으면 기존 동작):
    #   table_header_fill/table_header_text — 라이트 헤더(밝은 배경+진한 글자) 룩
    #   zebra(기본 True) — 짝수행 줄무늬
    hdr_fill = c.get("table_header_fill", c["primary"])
    hdr_text = c.get("table_header_text", c["header_font"])
    zebra = th.get("zebra", True)
    # 테두리 옵션(없으면 기존 박스형 동작):
    #   table_sides(기본 True) — 좌/우 외곽선 표시(False면 삼선표 느낌)
    #   table_edge_sz(기본 18=2.25pt) — 상/하·좌/우 외곽 굵기(1/8pt)
    #   table_header_sep(기본 False) — 헤더 하단·합계 상단을 옅은 가로 구분선으로
    sides = th.get("table_sides", True)
    edge_sz = th.get("table_edge_sz", 18)
    sep = th.get("table_header_sep", False)
    edge_col, in_col, in_sz = c["border_out"], c["border_in"], 4

    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    def border_for(ri, ci, total=False):
        ed = {}
        ed["top"] = (edge_sz, edge_col) if ri == 0 else (in_sz, in_col)
        ed["bottom"] = (edge_sz, edge_col) if ri == len(rows) else (in_sz, in_col)
        if ci == 0:
            if sides:
                ed["left"] = (edge_sz, edge_col)        # 좌 외곽(끄면 생략)
        else:
            ed["left"] = (in_sz, in_col)                # 열 사이 가는 선
        if ci == n_cols - 1:
            if sides:
                ed["right"] = (edge_sz, edge_col)       # 우 외곽(끄면 생략)
        else:
            ed["right"] = (in_sz, in_col)
        if ri == 0 and sep:                             # 헤더 하단 구분선
            ed["bottom"] = (8, c["secondary_lt"])
        if total:
            ed["top"] = (8, c["secondary_lt"]) if sep else (12, c["border_out"])
        return ed

    # 헤더
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        _shade_cell(cell, hdr_fill)
        _set_cell_borders(cell, border_for(0, ci))
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(h))
        _set_run_font(run, th["font_name"], size=s["header"], color=hdr_text, bold=True)

    # 본문
    for r, row in enumerate(rows, start=1):
        is_total = total_last_row and r == len(rows)
        for ci, val in enumerate(row):
            cell = tbl.cell(r, ci)
            # 줄무늬
            if zebra and not is_total and (r % 2 == 0):
                _shade_cell(cell, c["band"])
            if is_total:
                _shade_cell(cell, c["band"])
            _set_cell_borders(cell, border_for(r, ci, total=is_total))
            para = cell.paragraphs[0]
            is_num = ci in number_cols or isinstance(val, (int, float))
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_num else WD_ALIGN_PARAGRAPH.LEFT
            text = val
            if isinstance(val, (int, float)):
                text = _fmt_number(val, number_cols.get(ci))
            run = para.add_run(str(text))
            neg = isinstance(val, (int, float)) and val < 0
            _set_run_font(run, th["font_name"], size=s["body"],
                          color=c["negative"] if neg else c["text"], bold=is_total)

    if autofit:
        _enable_table_autofit(tbl, mode="window" if autofit == "window" else "content")
        for row in tbl.rows:
            for cell in row.cells:
                _tighten_cell(cell)
    return tbl


def add_page_number_footer(doc, theme=None):
    """바닥글 가운데에 페이지 번호 필드."""
    th = _resolve_theme(doc, theme)
    c, font = th["colors"], th["font_name"]
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = font
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb(c["secondary"])
    # PAGE 필드
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldEnd)
    return doc
