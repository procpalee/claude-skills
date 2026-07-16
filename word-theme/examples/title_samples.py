# -*- coding: utf-8 -*-
"""
제목부 스타일 6종 비교 샘플 생성기.

테마(default/procpa)마다 한 문서에 6개 스타일을 페이지별로 담아
어떤 제목부 디자인을 채택할지 눈으로 비교할 수 있게 한다.
    python title_samples.py
  → title_samples_default.docx, title_samples_procpa.docx
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx.enum.text import WD_BREAK
from word_theme import new_doc, add_title, add_heading, add_paragraph, add_table
from word_themes import WORD_THEMES

# (스타일, 라벨, 제목, 부제, meta) — 스타일별 대표 용례
# A 는 두 줄로 꺾이는 긴 제목으로 고정 줄높이(줄간격 압축)를 확인한다.
CASES = [
    ("bar", "시안 A — bar (기본: 제목 + 강조 바 · 두 줄 제목 줄간격 확인용)",
     "FY25 유형자산·사용권자산 손상평가 검토보고서 (회수가능액 산정 포함)",
     "OO사 · 2026. 3.", None),
    ("side", "시안 C — side (좌측 세로 악센트 바 · 간단한 메모/검토)",
     "회계처리 검토 메모", "리스 변경 회계처리 · 2026. 3.", None),
    ("band", "시안 D — band (전폭 색 밴드 · 제안서/트렌디 산출물)",
     "회계·재무 AX 도입 제안서", "OO 주식회사 · 2026. 3.",
     {"caption": "PROCPA CONSULTING"}),
    ("center", "시안 E — center (중앙 정렬 + 상하 가는 선 88% 폭 · 공식 의견서)",
     "회 신 문", "전환사채 발행자 회계처리 질의에 대한 의견",
     {"수신": "OO 주식회사", "문서번호": "26-031"}),
    ("meta", "시안 F — meta (제목 좌 + 문서정보 표 우 · 조서)",
     "수금조서 — 매출채권", "OO사 FY25 기말감사",
     {"조서번호": "C-210", "작성": "홍길동 3/15", "검토": "김철수 3/18"}),
]


def build(theme, out):
    doc = new_doc(theme)
    for i, (style, label, title, subtitle, meta) in enumerate(CASES):
        if i:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_paragraph(doc, label, color="9AA0AA")
        add_title(doc, title, subtitle, style=style, meta=meta)
        add_paragraph(doc, "본문이 여기서 시작됩니다. 제목부와 본문 사이 간격, "
                           "선 굵기, 색 대비를 확인하세요.")
        add_table(doc, ["계정과목", "전기", "당기"],
                  [["매출액", 1_200_000, 1_450_000],
                   ["영업이익", 220_000, 290_000]],
                  number_cols={1: "accounting", 2: "accounting"})
    doc.save(out)
    print("  생성:", out)


def main():
    here = os.path.dirname(os.path.abspath(__file__))
    print("제목부 샘플 생성 중...")
    for name in WORD_THEMES:
        build(name, os.path.join(here, f"title_samples_{name}.docx"))
    print("완료.")


if __name__ == "__main__":
    main()
